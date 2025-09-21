#!/usr/bin/env python3
"""
Script to extract text content from a .docx file
"""

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx library not found. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def extract_docx_content(docx_path):
    """
    Extract text content from a .docx file
    
    Args:
        docx_path (str): Path to the .docx file
    
    Returns:
        str: Extracted text content
    """
    try:
        # Load the document
        doc = Document(docx_path)
        
        # Extract text from all paragraphs
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text)
    
    except Exception as e:
        return f"Error reading document: {str(e)}"

def main():
    # Path to the .docx file
    docx_file = "Week 1 Personal Care Scenario -   Fall 2025.doc"
    
    if not Path(docx_file).exists():
        print(f"Error: File '{docx_file}' not found in current directory")
        return
    
    print(f"Extracting content from: {docx_file}")
    print("=" * 50)
    
    # Extract content
    content = extract_docx_content(docx_file)
    
    # Print content to console
    print(content)
    
    # Also save to a text file
    output_file = "extracted_content.txt"
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"Content extracted from: {docx_file}\n")
            f.write("=" * 50 + "\n\n")
            f.write(content)
        print(f"\n\nContent also saved to: {output_file}")
    except Exception as e:
        print(f"Error saving to file: {str(e)}")

if __name__ == "__main__":
    main()
